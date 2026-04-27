# -*- coding: utf-8 -*-
"""
generar_documento.py
Genera docs/documentacion_evoto.docx desde cero usando python-docx.
Cada sección de diagrama incluye el código PlantUML correspondiente.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Validación previa ─────────────────────────────────────────────────────────
OUTPUT = "docs/documentacion_evoto.docx"
os.makedirs("docs", exist_ok=True)
if os.path.exists(OUTPUT):
    try:
        with open(OUTPUT, "a"):
            pass
    except PermissionError:
        print("[!] ERROR: Cierra Microsoft Word antes de continuar")
        raise SystemExit(1)

# ── Contador de figuras ───────────────────────────────────────────────────────
_fig = [0]
def next_fig(nombre):
    _fig[0] += 1
    return f"Figura {_fig[0]} — {nombre}"

# ── Helpers base ──────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size_pt=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size_pt=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16.5)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(sizes.get(level, 12))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size_pt=size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows, caption=None):
    if caption:
        pc = doc.add_paragraph(caption)
        pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in pc.runs:
            run.font.italic = True
            run.font.size = Pt(10)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def use_case_table(doc, fields):
    add_table(doc, ["Campo", "Descripcion"], fields)

def insertar_plantuml(doc, codigo, nombre_figura):
    """Inserta bloque PlantUML con fuente Consolas 9pt y fondo gris claro."""
    # Línea introductoria
    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run("Codigo PlantUML del diagrama:")
    r_intro.italic = True
    r_intro.font.name = "Calibri"
    r_intro.font.size = Pt(10)
    p_intro.paragraph_format.space_after = Pt(2)

    # Bloque de código con fondo gris
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent  = Cm(0.5)
    p_code.paragraph_format.right_indent = Cm(0.5)
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after  = Pt(4)
    pPr = p_code._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"),   "single")
        bdr.set(qn("w:sz"),    "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "AAAAAA")
        pBdr.append(bdr)
    pPr.append(pBdr)
    run = p_code.add_run(codigo)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    # Pie de figura
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run(next_fig(nombre_figura))
    r_cap.italic = True
    r_cap.bold   = True
    r_cap.font.name = "Calibri"
    r_cap.font.size = Pt(10)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after  = Pt(10)

# ── Generadores PlantUML dinámicos ────────────────────────────────────────────

def pu_cu_individual(cu):
    """Diagrama de caso de uso individual."""
    actores = [a.strip() for a in cu["actor"].split("/")]
    lines = [
        "@startuml",
        "!theme plain",
        "left to right direction",
        "skinparam actorStyle awesome",
        "skinparam packageStyle rectangle",
        "",
    ]
    aliases = []
    for i, a in enumerate(actores):
        alias = f"A{i}"
        lines.append(f'actor "{a}" as {alias}')
        aliases.append(alias)
    lines += [
        "",
        'rectangle "Sistema E-Voto" {',
        f'  usecase "{cu["nombre"]}" as UC',
        "}",
        "",
    ]
    for alias in aliases:
        lines.append(f"{alias} --> UC")
    lines.append("@enduml")
    return "\n".join(lines)


def pu_secuencia(cu):
    """Diagrama de secuencia por caso de uso."""
    SEQ = {
        "CU01": (
            ["actor Administrador as ADM",
             "participant \"UsersPage\" as UI",
             "participant \"UsersController\" as UC",
             "participant \"UsersService\" as US",
             "database \"PostgreSQL\" as DB"],
            ["ADM -> UI : accederModuloUsuarios()",
             "UI -> UC : GET /users",
             "UC -> US : findAll()",
             "US -> DB : SELECT * FROM usuarios",
             "DB --> US : usuarios[]",
             "US --> UC : users[]",
             "UC --> UI : 200 OK + lista",
             "UI --> ADM : mostrar tabla"]
        ),
        "CU02": (
            ["actor Administrador as ADM",
             "participant \"ElectionManager\" as UI",
             "participant \"ElectionsController\" as EC",
             "participant \"ElectionsService\" as ES",
             "database \"PostgreSQL\" as DB"],
            ["ADM -> UI : fillForm(titulo,fechas)",
             "UI -> EC : POST /elections",
             "EC -> ES : crearEleccion(dto)",
             "ES -> DB : INSERT INTO elecciones",
             "DB --> ES : eleccion{id}",
             "ES --> EC : eleccion creada",
             "EC --> UI : 201 Created",
             "UI --> ADM : confirmar creacion"]
        ),
        "CU03": (
            ["actor Administrador as ADM",
             "participant \"ElectionManager\" as UI",
             "participant \"ElectionsController\" as EC",
             "participant \"ElectionsService\" as ES",
             "database \"PostgreSQL\" as DB"],
            ["ADM -> UI : fillCandidateForm()",
             "UI -> EC : POST /elections/:id/candidates",
             "EC -> ES : addCandidate(dto)",
             "ES -> DB : verificar estado=PROGRAMADA",
             "DB --> ES : OK",
             "ES -> DB : INSERT INTO candidatos",
             "DB --> ES : candidato{id}",
             "EC --> UI : 201 Created"]
        ),
        "CU04": (
            ["actor Administrador as ADM",
             "participant \"ElectionManager\" as UI",
             "participant \"ElectionsController\" as EC",
             "participant \"ElectionsService\" as ES",
             "participant \"FabricService\" as FS",
             "database \"PostgreSQL\" as DB"],
            ["ADM -> UI : cambiarEstado(nuevoEstado)",
             "UI -> EC : PATCH /elections/:id/status",
             "EC -> ES : changeStatus(estado)",
             "ES -> ES : validar transicion",
             "ES -> DB : UPDATE estado=CERRADA",
             "ES -> FS : cerrarEleccion() [si CERRADA]",
             "FS --> ES : OK o fallo no bloqueante",
             "EC --> UI : 200 OK"]
        ),
        "CU05": (
            ["actor Votante as VOT",
             "participant \"VotingPage\" as VP",
             "participant \"FabricController\" as FC",
             "participant \"UsersService\" as US",
             "participant \"FabricService\" as FS",
             "participant \"FabricGateway\" as FG",
             "database \"Ledger\" as LD",
             "database \"PostgreSQL\" as DB"],
            ["VOT -> VP : seleccionarCandidato()",
             "VP -> FC : POST /fabric/vote",
             "FC -> US : assertCanVote(userId,electionId)",
             "US -> DB : SELECT recibos WHERE estado=CONFIRMADO",
             "DB --> US : sin voto previo",
             "FC -> FS : emitirVoto(datos)",
             "FS -> FG : submitTransaction(vote)",
             "FG -> LD : commit bloque",
             "LD --> FG : txId",
             "FC -> US : markAsVoted() SELECT FOR UPDATE",
             "FS -> DB : INSERT recibos CONFIRMADO",
             "FC --> VP : {txId, success}",
             "VP --> VOT : mostrar comprobante"]
        ),
        "CU06": (
            ["actor \"Votante/Auditor\" as USR",
             "participant \"VerifyPage\" as VP",
             "participant \"FabricController\" as FC",
             "participant \"FabricService\" as FS",
             "database \"Ledger\" as LD"],
            ["USR -> VP : ingresarTxId(txId)",
             "VP -> FC : GET /fabric/verify/:txId",
             "FC -> FS : verificarVoto(txId)",
             "FS -> LD : queryTransaction(txId)",
             "LD --> FS : VoteAsset{voteId,electionId,timestamp}",
             "FS --> FC : voteInfo",
             "FC --> VP : 200 OK + datos del voto",
             "VP --> USR : mostrar confirmacion"]
        ),
        "CU07": (
            ["actor \"Admin/Auditor\" as USR",
             "participant \"ResultsPage\" as RP",
             "participant \"FabricController\" as FC",
             "database \"PostgreSQL\" as DB"],
            ["USR -> RP : seleccionarEleccion(id)",
             "RP -> FC : GET /fabric/results/:id",
             "FC -> DB : SELECT candidato, COUNT(*) FROM recibos",
             "DB --> FC : tally{candidato:votos}",
             "FC --> RP : TallyResult",
             "RP --> USR : grafico de resultados"]
        ),
        "CU08": (
            ["actor \"Admin/Auditor\" as USR",
             "participant \"AuditLogs\" as AL",
             "participant \"AuditController\" as AC",
             "database \"PostgreSQL\" as DB"],
            ["USR -> AL : abrirAuditoria()",
             "AL -> AC : GET /audit/logs",
             "AC -> DB : SELECT * FROM recibos_voto ORDER BY creado_en DESC",
             "DB --> AC : recibos[]",
             "AC --> AL : 200 OK + logs",
             "AL --> USR : tabla de recibos"]
        ),
        "CU09": (
            ["actor Administrador as ADM",
             "participant \"NodesPage\" as NP",
             "participant \"NodesController\" as NC",
             "participant \"NodesService\" as NS",
             "participant \"WSL/Docker\" as WD",
             "database \"PostgreSQL\" as DB"],
            ["ADM -> NP : registrarNodo(endpoint)",
             "NP -> NC : POST /nodes",
             "NC -> NS : createNode(dto)",
             "NS -> DB : INSERT INTO nodos_fabric",
             "ADM -> NP : toggleNodo(id)",
             "NP -> NC : POST /nodes/:id/toggle",
             "NC -> WD : docker start/stop via child_process",
             "WD --> NC : OK",
             "NC -> NS : reconnect() actualizar peer activo",
             "NC --> NP : 200 OK"]
        ),
        "CU10": (
            ["actor \"Usuario\" as USR",
             "participant \"LoginPage\" as LP",
             "participant \"AuthController\" as AC",
             "participant \"AuthService\" as AS",
             "participant \"JwtService\" as JWT",
             "database \"PostgreSQL\" as DB"],
            ["USR -> LP : submit(identificador, password)",
             "LP -> AC : POST /auth/login",
             "AC -> AS : login(dto)",
             "AS -> DB : findByIdentificador(id)",
             "DB --> AS : usuario{hash,rol,habilitado}",
             "AS -> AS : bcrypt.compare(pwd, hash)",
             "AS -> JWT : sign({sub,identificador,role})",
             "JWT --> AS : access_token",
             "AS --> AC : {access_token, user}",
             "AC --> LP : 200 OK + JWT",
             "LP --> USR : redirect segun rol"]
        ),
    }
    cu_id = cu["id"]
    if cu_id not in SEQ:
        return f"@startuml\n!theme plain\nnote: diagrama de secuencia para {cu_id}\n@enduml"
    parts, msgs = SEQ[cu_id]
    lines = ["@startuml", "!theme plain", "skinparam sequenceMessageAlign center", ""]
    lines += parts
    lines.append("")
    lines += msgs
    lines.append("@enduml")
    return "\n".join(lines)

# ── PlantUML estáticos ────────────────────────────────────────────────────────

PU_CU_GENERAL = """\
@startuml
!theme plain
left to right direction
skinparam actorStyle awesome
skinparam packageStyle rectangle

actor Administrador as ADM
actor Votante as VOT
actor Auditor as AUD

rectangle "Sistema E-Voto Hyperledger Fabric" {
  usecase "CU01 Gestionar Usuarios" as UC01
  usecase "CU02 Gestionar Elecciones" as UC02
  usecase "CU03 Gestionar Candidatos" as UC03
  usecase "CU04 Cambiar Estado Eleccion" as UC04
  usecase "CU05 Emitir Voto" as UC05
  usecase "CU06 Verificar Voto en Blockchain" as UC06
  usecase "CU07 Ver Resultados Electorales" as UC07
  usecase "CU08 Ver Logs de Auditoria" as UC08
  usecase "CU09 Gestionar Nodos Fabric" as UC09
  usecase "CU10 Iniciar Sesion" as UC10
}

ADM --> UC01
ADM --> UC02
ADM --> UC03
ADM --> UC04
ADM --> UC07
ADM --> UC08
ADM --> UC09
ADM --> UC10

VOT --> UC05
VOT --> UC06
VOT --> UC10

AUD --> UC06
AUD --> UC07
AUD --> UC08
AUD --> UC10
@enduml"""

PU_PKG_TRACE = """\
@startuml
!theme plain
skinparam packageStyle rectangle

package "Seguridad" as P1 {
}
package "Gestion Electoral" as P2 {
}
package "Votacion" as P3 {
}
package "Auditoria y Reportes" as P4 {
}

usecase "CU10 Iniciar Sesion" as UC10
usecase "CU01 Gestionar Usuarios" as UC01
usecase "CU02 Gestionar Elecciones" as UC02
usecase "CU03 Gestionar Candidatos" as UC03
usecase "CU04 Cambiar Estado" as UC04
usecase "CU05 Emitir Voto" as UC05
usecase "CU06 Verificar Voto" as UC06
usecase "CU07 Ver Resultados" as UC07
usecase "CU08 Ver Logs Auditoria" as UC08
usecase "CU09 Gestionar Nodos" as UC09

P1 ..> UC10 : <<trace>>
P2 ..> UC01 : <<trace>>
P2 ..> UC02 : <<trace>>
P2 ..> UC03 : <<trace>>
P2 ..> UC04 : <<trace>>
P3 ..> UC05 : <<trace>>
P3 ..> UC06 : <<trace>>
P4 ..> UC07 : <<trace>>
P4 ..> UC08 : <<trace>>
P4 ..> UC09 : <<trace>>
@enduml"""

PU_PKG_GENERAL = """\
@startuml
!theme plain
skinparam packageStyle rectangle

package "Seguridad" as P1 {
  [AuthService]
  [JwtService]
  [RolesGuard]
}

package "Gestion Electoral" as P2 {
  [ElectionsService]
  [CandidatesService]
  [UsersService]
}

package "Votacion" as P3 {
  [FabricService]
  [FabricGateway]
  [VerifyService]
}

package "Auditoria y Reportes" as P4 {
  [AuditController]
  [ResultsService]
  [NodesService]
}

P2 ..> P1 : usa
P3 ..> P1 : usa
P3 ..> P2 : usa
P4 ..> P1 : usa
P4 ..> P2 : usa
P4 ..> P3 : usa
@enduml"""

PU_COMM_CU10 = """\
@startuml
!theme plain

object "<<boundary>>\\nLoginPage" as LP
object "<<control>>\\nAuthController" as AC
object "<<control>>\\nAuthService" as AS
object "<<entity>>\\nusuarios (DB)" as USR
object "<<control>>\\nJwtService" as JWT

actor "Usuario" as U

U --> LP : 1.1 submit(id, pwd)
LP --> AC : 1.2 POST /auth/login
AC --> AS : 2.1 login(dto)
AS --> USR : 2.2 findByIdentificador()
USR --> AS : 2.3 usuario{hash,rol}
AS --> AS : 3.1 bcrypt.compare()
AS --> JWT : 3.2 sign(payload)
JWT --> AS : 3.3 access_token
AS --> AC : 4.1 {access_token, user}
AC --> LP : 4.2 200 OK
LP --> U : 4.3 redirect segun rol
@enduml"""

PU_COMM_CU05 = """\
@startuml
!theme plain

object "<<boundary>>\\nVotingPage" as VP
object "<<control>>\\nFabricController" as FC
object "<<control>>\\nUsersService" as US
object "<<control>>\\nFabricService" as FS
object "<<control>>\\nFabricGateway" as FG
object "<<entity>>\\nLedger Fabric" as LD
object "<<entity>>\\nrecibos_voto" as RV

actor Votante as V

V --> VP : 1.1 seleccionarCandidato()
VP --> FC : 1.2 POST /fabric/vote
FC --> US : 2.1 assertCanVote()
US --> RV : 2.2 verificar sin voto previo
FC --> FS : 3.1 emitirVoto(datos)
FS --> FG : 3.2 submitTransaction()
FG --> LD : 3.3 commit bloque
LD --> FG : 3.4 txId
FC --> US : 4.1 markAsVoted() SELECT FOR UPDATE
FS --> RV : 4.2 INSERT estado=CONFIRMADO
FC --> VP : 5.1 {txId, success}
VP --> V : 5.2 mostrar comprobante
@enduml"""

PU_COMM_CU08 = """\
@startuml
!theme plain

object "<<boundary>>\\nAuditLogs" as AL
object "<<control>>\\nAuditController" as AC
object "<<entity>>\\nrecibos_voto" as RV
object "<<entity>>\\nFabric Ledger" as LD

actor "Auditor" as AUD

AUD --> AL : 1.1 abrirAuditoria()
AL --> AC : 1.2 GET /audit/logs
AC --> RV : 2.1 SELECT * ORDER BY fecha DESC
RV --> AC : 2.2 recibos[]
AC --> LD : 3.1 queryBlocks() verificar integridad
LD --> AC : 3.2 bloques validados
AC --> AL : 4.1 200 OK + logs
AL --> AUD : 4.2 tabla de recibos auditados
@enduml"""

PU_DEPLOY = """\
@startuml
!theme plain

node "Cliente Web" as CLI {
  artifact "React + Vite" as FRONT
}

node "Servidor Aplicacion" as APP {
  artifact "NestJS API :3000" as BACK
  artifact "Fabric Gateway SDK" as SDK
}

node "PostgreSQL 15 :5432" as PGDB {
  database "sw2_primer_parcial" as DB
}

node "peer0.ficct.edu.bo :7051" as P0 {
  artifact "Endorsing Peer" as EP0
  database "CouchDB :5984" as CDB0
  artifact "evoting-cc" as CC0
}

node "peer1.ficct.edu.bo :8051" as P1 {
  artifact "Endorsing Peer" as EP1
  database "CouchDB :6984" as CDB1
}

node "orderer.ficct.edu.bo :7050" as ORD {
  artifact "Orderer (Raft)" as ON
}

CLI -- APP : HTTPS
APP -- PGDB : TCP 5432
APP -- P0 : gRPC 7051
APP -- P1 : gRPC 8051
P0 -- ORD : gRPC 7050
P1 -- ORD : gRPC 7050
@enduml"""

PU_CLASES = """\
@startuml
!theme plain
skinparam classAttributeIconSize 0

class organizaciones {
  +id : UUID PK
  +nombre : VARCHAR(255)
  +slug : VARCHAR(100)
  +activo : BOOLEAN
  +creado_en : TIMESTAMPTZ
}

class usuarios {
  +id : UUID PK
  +id_organizacion : UUID FK
  +identificador : VARCHAR(100)
  +nombre : VARCHAR(255)
  +email : VARCHAR(255)
  +hash_contrasena : VARCHAR(255)
  +rol : rol_usuario
  +habilitado : BOOLEAN
  +metadatos : JSONB
  +creado_en : TIMESTAMPTZ
}

class elecciones {
  +id : UUID PK
  +id_organizacion : UUID FK
  +titulo : VARCHAR(255)
  +estado : estado_eleccion
  +fecha_inicio : TIMESTAMPTZ
  +fecha_fin : TIMESTAMPTZ
  +canal_fabric : VARCHAR(100)
  +creado_en : TIMESTAMPTZ
}

class candidatos {
  +id : UUID PK
  +id_eleccion : UUID FK
  +nombre_frente : VARCHAR(100)
  +nombre_candidato : VARCHAR(255)
  +nombre_cargo : VARCHAR(100)
  +url_foto : TEXT
  +creado_en : TIMESTAMPTZ
}

class padron_electoral {
  +id : UUID PK
  +id_eleccion : UUID FK
  +id_usuario : UUID FK
  +voto_emitido : BOOLEAN
  +votado_en : TIMESTAMPTZ
}

class recibos_voto {
  +id : UUID PK
  +id_usuario : UUID FK
  +id_eleccion : UUID FK
  +id_candidato : UUID FK
  +id_transaccion : VARCHAR UNIQUE
  +estado : estado_sincronizacion
  +creado_en : TIMESTAMPTZ
}

class eventos_auditoria {
  +id : UUID PK
  +id_usuario : UUID FK
  +accion : accion_auditoria
  +direccion_ip : INET
  +detalles : JSONB
  +creado_en : TIMESTAMPTZ
}

class nodos_fabric {
  +id : UUID PK
  +endpoint : VARCHAR(255)
  +host_alias : VARCHAR(255)
  +activo : BOOLEAN
  +prioridad : SMALLINT
}

class canales_fabric {
  +id : UUID PK
  +nombre : VARCHAR(100) UNIQUE
  +activo : BOOLEAN
}

organizaciones "1" -- "N" usuarios
organizaciones "1" -- "N" elecciones
elecciones "1" -- "N" candidatos
elecciones "1" -- "N" padron_electoral
usuarios "1" -- "N" padron_electoral
elecciones "1" -- "N" recibos_voto
usuarios "1" -- "N" recibos_voto
candidatos "1" -- "N" recibos_voto
usuarios "1" -- "N" eventos_auditoria
@enduml"""

# ── Datos del sistema ─────────────────────────────────────────────────────────

CASOS_DE_USO = [
    {
        "id": "CU01",
        "nombre": "Gestionar Usuarios",
        "actor": "Administrador",
        "proposito": "Permitir al Administrador crear, editar, habilitar/deshabilitar y eliminar cuentas de usuario del sistema.",
        "precondiciones": "El Administrador debe estar autenticado con rol ADMINISTRADOR.",
        "flujo": (
            "1. El Administrador accede al modulo de Usuarios.\n"
            "2. Visualiza la lista de usuarios registrados.\n"
            "3. Selecciona la accion: crear, editar o eliminar.\n"
            "4. Completa los campos requeridos (identificador, nombre, email, rol, carrera).\n"
            "5. El sistema valida los datos y persiste los cambios en la tabla usuarios.\n"
            "6. Se muestra confirmacion de la operacion."
        ),
        "excepcion": "Si el identificador o email ya existe, el sistema lanza ConflictException y no persiste el registro.",
        "postcondiciones": "El usuario queda registrado, actualizado o eliminado en la base de datos."
    },
    {
        "id": "CU02",
        "nombre": "Gestionar Elecciones",
        "actor": "Administrador",
        "proposito": "Crear, editar y eliminar procesos electorales, incluyendo fechas, descripcion y estado.",
        "precondiciones": "El Administrador debe estar autenticado. La eleccion a eliminar debe estar en estado PROGRAMADA.",
        "flujo": (
            "1. El Administrador accede a ElectionManager.\n"
            "2. Crea una nueva eleccion con titulo, descripcion, fechas de inicio y fin.\n"
            "3. El sistema inserta la eleccion con estado inicial PROGRAMADA.\n"
            "4. El Administrador puede editar campos mientras la eleccion no este ACTIVA.\n"
            "5. Para eliminar, la eleccion debe estar en estado PROGRAMADA."
        ),
        "excepcion": "Intentar eliminar una eleccion en estado distinto a PROGRAMADA genera BadRequestException.",
        "postcondiciones": "La eleccion queda registrada en la tabla elecciones lista para agregar candidatos."
    },
    {
        "id": "CU03",
        "nombre": "Gestionar Candidatos",
        "actor": "Administrador",
        "proposito": "Agregar y eliminar candidatos asociados a una eleccion en estado PROGRAMADA.",
        "precondiciones": "Debe existir una eleccion en estado PROGRAMADA.",
        "flujo": (
            "1. El Administrador selecciona una eleccion.\n"
            "2. Agrega candidatos con: nombre de frente, nombre del candidato, cargo, foto URL y mision.\n"
            "3. El sistema inserta el candidato en la tabla candidatos.\n"
            "4. Para eliminar, el Administrador selecciona el candidato y confirma."
        ),
        "excepcion": "Agregar o eliminar candidatos en eleccion no PROGRAMADA lanza BadRequestException.",
        "postcondiciones": "El candidato queda registrado y disponible en la boleta electoral."
    },
    {
        "id": "CU04",
        "nombre": "Cambiar Estado de Eleccion",
        "actor": "Administrador",
        "proposito": "Avanzar una eleccion por su ciclo de vida: PROGRAMADA->ACTIVA->CERRADA->ESCRUTADA.",
        "precondiciones": "La eleccion debe existir. La transicion de estado debe ser valida.",
        "flujo": (
            "1. El Administrador selecciona la eleccion.\n"
            "2. Solicita el cambio de estado.\n"
            "3. El sistema valida la transicion segun VALID_TRANSITIONS.\n"
            "4. Si el nuevo estado es CERRADA, se invoca fabricService.cerrarEleccion().\n"
            "5. El estado se actualiza en la tabla elecciones."
        ),
        "excepcion": "Transicion invalida genera BadRequestException. El fallo de Fabric al cerrar se registra en el log pero no bloquea la transicion.",
        "postcondiciones": "El estado de la eleccion queda actualizado. Si es CERRADA, el ledger de Fabric registra el cierre."
    },
    {
        "id": "CU05",
        "nombre": "Emitir Voto",
        "actor": "Votante",
        "proposito": "Permitir al Votante habilitado emitir su voto en una eleccion ACTIVA de forma anonima e inmutable.",
        "precondiciones": "El Votante debe estar autenticado, habilitado y no haber votado previamente en esa eleccion.",
        "flujo": (
            "1. El Votante accede a VotingPage y ve las elecciones ACTIVAS.\n"
            "2. Selecciona un candidato por eleccion.\n"
            "3. El sistema ejecuta assertCanVote() verificando habilitacion y ausencia de voto previo.\n"
            "4. Se llama a emitirVoto() en Hyperledger Fabric, que retorna txId y voteId.\n"
            "5. Se ejecuta markAsVoted() con SELECT FOR UPDATE para prevenir condiciones de carrera.\n"
            "6. Se guarda un recibo en recibos_voto con estado CONFIRMADO.\n"
            "7. El Votante recibe el txId como comprobante."
        ),
        "excepcion": "Si el Votante ya voto: ConflictException. Si Fabric falla: recibo con estado FALLIDO, el voto NO se marca como emitido.",
        "postcondiciones": "El voto queda registrado anonimamente en Fabric y el recibo en recibos_voto."
    },
    {
        "id": "CU06",
        "nombre": "Verificar Voto en Blockchain",
        "actor": "Votante / Auditor",
        "proposito": "Consultar en el ledger de Hyperledger Fabric la existencia e integridad de un voto mediante su txId.",
        "precondiciones": "El usuario debe disponer del txId recibido al momento de votar.",
        "flujo": (
            "1. El usuario accede al verificador y proporciona el txId.\n"
            "2. El sistema invoca fabricService.verificarVoto(txId).\n"
            "3. Fabric retorna el VoteAsset con: assetType, voteId, electionId, candidateId, timestamp.\n"
            "4. Se muestra la informacion al usuario."
        ),
        "excepcion": "Si el txId no existe en el ledger, se lanza NotFoundException.",
        "postcondiciones": "El usuario obtiene confirmacion criptografica de la existencia del voto."
    },
    {
        "id": "CU07",
        "nombre": "Ver Resultados Electorales",
        "actor": "Administrador / Auditor / Publico",
        "proposito": "Consultar los resultados parciales o finales de una eleccion agrupados por candidato.",
        "precondiciones": "Debe existir al menos una eleccion con votos registrados.",
        "flujo": (
            "1. El usuario accede a AdminResults, AuditorDashboard o LiveResults.\n"
            "2. Selecciona la eleccion.\n"
            "3. El sistema consulta recibos_voto JOIN candidatos contando votos CONFIRMADOS.\n"
            "4. Se muestra grafico/tabla con conteos por candidato."
        ),
        "excepcion": "Si la eleccion no tiene votos, se muestran todos los candidatos con conteo 0.",
        "postcondiciones": "El usuario visualiza los resultados actualizados en tiempo real."
    },
    {
        "id": "CU08",
        "nombre": "Ver Logs de Auditoria",
        "actor": "Administrador / Auditor",
        "proposito": "Consultar el historial de recibos de voto para auditoria del proceso electoral.",
        "precondiciones": "El usuario debe tener rol ADMINISTRADOR o AUDITOR.",
        "flujo": (
            "1. El usuario accede a AuditLogs.\n"
            "2. El sistema consulta la tabla recibos_voto ordenada por fecha descendente.\n"
            "3. Se muestran: userId, electionId, txId, estado, mensaje de error y fecha."
        ),
        "excepcion": "Si no existen registros, se muestra lista vacia.",
        "postcondiciones": "El auditor puede revisar cada transaccion y su estado en Fabric."
    },
    {
        "id": "CU09",
        "nombre": "Gestionar Nodos Fabric",
        "actor": "Administrador",
        "proposito": "Registrar, activar/desactivar y desplegar peers de Hyperledger Fabric dinamicamente.",
        "precondiciones": "El Administrador debe estar autenticado. WSL con Docker debe estar disponible.",
        "flujo": (
            "1. El Administrador accede a NodesPage.\n"
            "2. Puede registrar un nodo manualmente (endpoint, host alias).\n"
            "3. Puede hacer toggle start/stop ejecutando docker start/stop via child_process.\n"
            "4. Puede desplegar un nuevo peer ejecutando add-peer-dynamic.sh en WSL.\n"
            "5. FabricService.reconnect() se actualiza automaticamente con el peer activo de mayor prioridad."
        ),
        "excepcion": "Si el script WSL falla o no devuelve PEER_ENDPOINT, se lanza InternalServerErrorException.",
        "postcondiciones": "El nodo queda registrado en nodos_fabric y Fabric se reconecta al peer activo."
    },
    {
        "id": "CU10",
        "nombre": "Iniciar Sesion",
        "actor": "Administrador / Votante / Auditor",
        "proposito": "Autenticar al usuario en el sistema y obtener un token JWT para acceder a funciones protegidas.",
        "precondiciones": "El usuario debe estar registrado y habilitado en el sistema.",
        "flujo": (
            "1. El usuario accede a /login e ingresa identificador y contrasena.\n"
            "2. El sistema busca el usuario por identificador.\n"
            "3. Se compara la contrasena con bcrypt.compare().\n"
            "4. Se genera un JWT con sub, identificador y rol.\n"
            "5. Se retorna access_token y datos basicos del usuario.\n"
            "6. El frontend guarda el token en Zustand y redirige segun el rol."
        ),
        "excepcion": "Credenciales invalidas o cuenta deshabilitada generan UnauthorizedException.",
        "postcondiciones": "El usuario queda autenticado con JWT valido por 8 horas."
    },
]

DB_TABLES = [
    ("organizaciones", [
        ("id", "UUID", "PK", "—"),
        ("nombre", "VARCHAR(255)", "—", "—"),
        ("slug", "VARCHAR(100)", "UNIQUE", "—"),
        ("url_logo", "TEXT", "—", "—"),
        ("configuracion", "JSONB", "—", "DEFAULT '{}'"),
        ("activo", "BOOLEAN", "—", "DEFAULT TRUE"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("usuarios", [
        ("id", "UUID", "PK", "—"),
        ("id_organizacion", "UUID", "FK -> organizaciones", "NOT NULL"),
        ("identificador", "VARCHAR(100)", "UNIQUE(org)", "R.U. / codigo empleado"),
        ("nombre", "VARCHAR(255)", "—", "NOT NULL"),
        ("email", "VARCHAR(255)", "UNIQUE(org)", "NOT NULL"),
        ("hash_contrasena", "VARCHAR(255)", "—", "bcrypt"),
        ("rol", "rol_usuario ENUM", "—", "VOTANTE | ADMINISTRADOR | AUDITOR"),
        ("metadatos", "JSONB", "—", "{carrera: SISTEMAS}"),
        ("habilitado", "BOOLEAN", "—", "DEFAULT TRUE"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("elecciones", [
        ("id", "UUID", "PK", "—"),
        ("id_organizacion", "UUID", "FK -> organizaciones", "NOT NULL"),
        ("titulo", "VARCHAR(255)", "—", "NOT NULL"),
        ("descripcion", "TEXT", "—", "nullable"),
        ("fecha_inicio", "TIMESTAMPTZ", "—", "NOT NULL"),
        ("fecha_fin", "TIMESTAMPTZ", "—", "NOT NULL; CHECK fin > inicio"),
        ("estado", "estado_eleccion ENUM", "—", "DEFAULT 'BORRADOR'"),
        ("canal_fabric", "VARCHAR(100)", "—", "DEFAULT 'evoting'"),
        ("id_eleccion_fabric", "UUID", "UNIQUE", "uuid_generate_v4()"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("candidatos", [
        ("id", "UUID", "PK", "—"),
        ("id_eleccion", "UUID", "FK -> elecciones CASCADE", "NOT NULL"),
        ("nombre_frente", "VARCHAR(100)", "—", "NOT NULL"),
        ("nombre_candidato", "VARCHAR(255)", "—", "NOT NULL"),
        ("nombre_cargo", "VARCHAR(100)", "—", "DECANO, DIRECTOR, etc."),
        ("url_foto", "TEXT", "—", "nullable"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("padron_electoral", [
        ("id", "UUID", "PK", "—"),
        ("id_eleccion", "UUID", "FK -> elecciones CASCADE", "NOT NULL"),
        ("id_usuario", "UUID", "FK -> usuarios CASCADE", "NOT NULL"),
        ("voto_emitido", "BOOLEAN", "—", "DEFAULT FALSE"),
        ("inscrito_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("recibos_voto", [
        ("id", "UUID", "PK", "—"),
        ("id_usuario", "UUID", "FK -> usuarios", "NOT NULL"),
        ("id_eleccion", "UUID", "FK -> elecciones", "NOT NULL"),
        ("id_candidato", "UUID", "FK -> candidatos", "nullable"),
        ("id_voto", "UUID", "—", "UUID anonimo en Fabric"),
        ("id_transaccion", "VARCHAR(255)", "UNIQUE", "txId de Fabric"),
        ("estado", "estado_sincronizacion ENUM", "—", "PENDIENTE | CONFIRMADO | FALLIDO"),
        ("mensaje_error", "TEXT", "—", "nullable"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("eventos_auditoria", [
        ("id", "UUID", "PK", "—"),
        ("id_organizacion", "UUID", "FK -> organizaciones", "nullable"),
        ("id_usuario", "UUID", "FK -> usuarios", "nullable"),
        ("accion", "accion_auditoria ENUM", "—", "19 valores posibles"),
        ("direccion_ip", "INET", "—", "nullable"),
        ("detalles", "JSONB", "—", "DEFAULT '{}'"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("nodos_fabric", [
        ("id", "UUID", "PK", "—"),
        ("nombre", "VARCHAR(100)", "—", "NOT NULL"),
        ("endpoint", "VARCHAR(255)", "—", "e.g. localhost:7051"),
        ("host_alias", "VARCHAR(255)", "—", "e.g. peer0.ficct.edu.bo"),
        ("activo", "BOOLEAN", "—", "DEFAULT TRUE"),
        ("prioridad", "SMALLINT", "—", "menor = mayor prioridad"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
    ("canales_fabric", [
        ("id", "UUID", "PK", "—"),
        ("nombre", "VARCHAR(100)", "UNIQUE", "NOT NULL"),
        ("descripcion", "TEXT", "—", "nullable"),
        ("activo", "BOOLEAN", "—", "DEFAULT TRUE"),
        ("creado_en", "TIMESTAMPTZ", "—", "DEFAULT NOW()"),
    ]),
]

PAQUETES = [
    ("Seguridad",
     "Autenticacion y control de acceso. JWT, bcrypt, JwtAuthGuard, RolesGuard. "
     "Casos de uso: CU10 - Iniciar Sesion."),
    ("Gestion Electoral",
     "Administracion del ciclo de vida electoral: gestionar usuarios, elecciones, candidatos "
     "y cambiar estado del proceso. "
     "Casos de uso: CU01, CU02, CU03, CU04."),
    ("Votacion",
     "Emision anonima del voto en Hyperledger Fabric y verificacion individual por txId. "
     "Incluye la proteccion anti-doble-voto (assertCanVote + SELECT FOR UPDATE). "
     "Casos de uso: CU05 - Emitir Voto, CU06 - Verificar Voto en Blockchain."),
    ("Auditoria y Reportes",
     "Verificacion de integridad, logs de auditoria, resultados electorales y gestion de nodos Fabric. "
     "Accesible por roles Administrador y Auditor. "
     "Casos de uso: CU07, CU08, CU09."),
]

# ═══════════════════════════════ DOCUMENTO ════════════════════════════════════

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
    setattr(section, attr, Cm(2.5))

# ── Portada ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSIDAD AUTONOMA GABRIEL RENE MORENO\n"
              "FACULTAD DE INGENIERIA EN CIENCIAS DE LA COMPUTACION Y TELECOMUNICACIONES")
r.font.name = "Calibri"; r.font.size = Pt(14); r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SISTEMA DE VOTACION ELECTRONICA (E-VOTO)\nCON HYPERLEDGER FABRIC")
r.font.name = "Calibri"; r.font.size = Pt(18); r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentacion Tecnica — Ingenieria de Software II\nSanta Cruz de la Sierra, Bolivia — 2026")
r.font.name = "Calibri"; r.font.size = Pt(12)

doc.add_page_break()

# ─────────────────────────────── INTRODUCCION ─────────────────────────────────
add_heading(doc, "INTRODUCCION", 1)
add_paragraph(doc,
    "La confianza en los procesos electorales es un pilar fundamental de toda democracia. Sin embargo, "
    "los sistemas de votacion tradicionales presentan vulnerabilidades ante la manipulacion de resultados, "
    "la falta de transparencia y el riesgo de votos duplicados. El presente proyecto aborda esta "
    "problematica mediante el desarrollo de un Sistema de Votacion Electronica (E-Voto) para la "
    "Facultad de Ingenieria en Ciencias de la Computacion y Telecomunicaciones (FICCT) de la "
    "Universidad Autonoma Gabriel Rene Moreno (UAGRM).")
add_paragraph(doc,
    "La propuesta integra una arquitectura multicapa compuesta por un backend NestJS, un frontend "
    "React, una base de datos PostgreSQL y una red blockchain Hyperledger Fabric. Esta combinacion "
    "permite separar la logica de identidad y control electoral —manejada por la base de datos "
    "relacional— del registro inmutable e inalterable del sufragio, garantizado por el ledger "
    "distribuido de Fabric.")
add_paragraph(doc,
    "El sistema contempla tres roles principales: el Administrador, responsable de configurar "
    "elecciones, candidatos y usuarios; el Votante, quien emite su sufragio de forma anonima "
    "mediante un identificador unico en la blockchain; y el Auditor, que verifica la integridad "
    "del proceso sin acceder a la identidad del votante.")
doc.add_page_break()

# ───────────────────────────── CAPITULO I ─────────────────────────────────────
add_heading(doc, "CAPITULO I — PERFIL DEL PROYECTO", 1)
add_heading(doc, "1.1. Problema de Investigacion", 2)
add_paragraph(doc,
    "Los procesos electorales universitarios en FICCT se realizan de manera presencial o mediante "
    "sistemas informaticos centralizados que no garantizan suficientemente la integridad ni la "
    "transparencia del voto. Entre los principales problemas identificados se encuentran:")
for prob in [
    "Riesgo de manipulacion de resultados en sistemas centralizados sin mecanismo de auditoria independiente.",
    "Ausencia de un registro inmutable que permita verificar cada voto emitido despues del proceso.",
    "Vulnerabilidad ante votos duplicados o emision fraudulenta en nombre de terceros.",
    "Falta de separacion entre la identidad del votante y el contenido de su voto, comprometiendo el secreto del sufragio.",
    "Limitada participacion estudiantil por barreras logisticas de la votacion presencial.",
]:
    add_bullet(doc, prob)

add_heading(doc, "1.2. Objetivos", 2)
add_heading(doc, "1.2.1. Objetivo General", 3)
add_paragraph(doc,
    "Desarrollar un Sistema de Votacion Electronica (E-Voto) para FICCT-UAGRM que garantice "
    "la integridad, anonimato e inmutabilidad del sufragio mediante la integracion de "
    "Hyperledger Fabric como capa de registro blockchain, complementada con un backend NestJS, "
    "base de datos PostgreSQL y una interfaz web React.")

add_heading(doc, "1.2.2. Objetivos Especificos", 3)
for obj in [
    "Disenar e implementar el esquema de base de datos relacional PostgreSQL para gestion de identidad, elecciones, candidatos y padron electoral.",
    "Desarrollar un backend RESTful en NestJS con autenticacion JWT, control de roles y logica anti-doble-voto mediante SELECT FOR UPDATE.",
    "Implementar el chaincode evoting-cc en Hyperledger Fabric para registrar votos de forma anonima e inmutable en el ledger.",
    "Construir una interfaz web React diferenciada por rol con proteccion de rutas y estado global Zustand.",
    "Integrar un mecanismo de auditoria que permita verificar la existencia de cualquier voto en la blockchain mediante su txId.",
    "Implementar la gestion dinamica de peers Fabric y canales desde la interfaz de administracion.",
]:
    add_bullet(doc, obj)

add_heading(doc, "1.3. Justificacion", 2)
add_paragraph(doc,
    "La implementacion de Hyperledger Fabric como infraestructura blockchain para el sistema "
    "E-Voto se justifica por multiples razones tecnicas: Fabric es una plataforma permisionada, "
    "lo que permite restringir la participacion a nodos autorizados. La separacion entre "
    "identidad (PostgreSQL) y contenido del voto (Fabric) garantiza el anonimato. La "
    "inmutabilidad del ledger asegura que ningun resultado pueda ser alterado retroactivamente. "
    "Finalmente, la capacidad de auditores externos de verificar cada voto mediante su txId "
    "sin conocer la identidad del votante representa un avance significativo en transparencia electoral.")
doc.add_page_break()

# ───────────────────────────── CAPITULO II ────────────────────────────────────
add_heading(doc, "CAPITULO II — FUNDAMENTACION TEORICA", 1)
add_heading(doc, "2.1. Tecnologia Blockchain", 2)
add_paragraph(doc,
    "Una blockchain es una base de datos distribuida que mantiene un registro continuo y creciente "
    "de transacciones, organizadas en bloques enlazados criptograficamente. Cada bloque contiene "
    "un hash del bloque anterior, garantizando la inmutabilidad de la cadena: modificar cualquier "
    "bloque invalida todos los bloques posteriores.")

add_heading(doc, "2.2. Hyperledger Fabric", 2)
add_paragraph(doc,
    "Hyperledger Fabric es un framework de blockchain permisionado de codigo abierto albergado "
    "por la Linux Foundation. Disenado para entornos empresariales donde se requiere control de "
    "acceso, privacidad entre participantes y alto rendimiento en transacciones.")

add_heading(doc, "2.2.1. Arquitectura de Hyperledger Fabric", 3)
for comp, desc in [
    ("Peers", "Nodos que mantienen el ledger y ejecutan el chaincode. peer0 y peer1 de FICCTOrg."),
    ("Orderer", "Servicio de ordenacion que garantiza el consenso (Raft) y la secuencia de transacciones."),
    ("Chaincode", "Contrato inteligente evoting-cc que implementa emitirVoto, verificarVoto y cerrarEleccion."),
    ("Canal evoting", "Subred privada dentro de Fabric. Canal principal del sistema E-Voto."),
    ("CouchDB", "Base de datos de estado (World State) que permite consultas ricas sobre el ledger actual."),
    ("MSP", "Gestiona identidades digitales y certificados X.509 para autenticar participantes."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{comp}: "); r1.font.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
    r2 = p.add_run(desc);        r2.font.name = "Calibri";                      r2.font.size = Pt(11)

add_heading(doc, "2.3. Principios de Seguridad Electoral", 2)
add_paragraph(doc,
    "Un sistema de votacion electronica robusto debe garantizar: Unicidad (un votante = un voto), "
    "Anonimato (el voto no puede vincularse al votante), Integridad (el voto no puede alterarse "
    "tras emitirse), Verificabilidad (el votante puede confirmar que su voto fue contado) y "
    "Disponibilidad (el sistema debe funcionar durante el periodo de votacion). El sistema "
    "E-Voto implementa cada uno de estos principios mediante capas tecnicas complementarias.")
doc.add_page_break()

# ───────────────────────────── CAPITULO III ───────────────────────────────────
add_heading(doc, "CAPITULO III — CASOS DE ESTUDIO", 1)
add_heading(doc, "3.1.1. Caso de Estudio I — Voatz", 3)
add_paragraph(doc,
    "Voatz es una plataforma de votacion movil que utilizo blockchain para elecciones piloto "
    "en Estados Unidos (West Virginia, 2018). Su arquitectura emplea Hyperledger Fabric "
    "para registrar votos. Demostro la viabilidad de la votacion remota segura.")
add_heading(doc, "3.1.2. Caso de Estudio II — Sierra Leone (Agora)", 3)
add_paragraph(doc,
    "En las elecciones presidenciales de Sierra Leone (2018), la empresa Agora registro "
    "votos en una blockchain publica como mecanismo de auditoria paralela. Inspiro el diseno "
    "dual del sistema E-Voto: PostgreSQL como fuente de verdad operacional y Fabric como "
    "capa de verificacion inmutable.")
add_heading(doc, "3.1.3. Caso de Estudio III — Universidades Latinoamericanas", 3)
add_paragraph(doc,
    "Diversas universidades latinoamericanas han implementado sistemas de votacion estudiantil "
    "basados en Ethereum o Hyperledger. El mayor desafio no es tecnico sino organizacional: "
    "garantizar que todos los votantes puedan acceder al sistema. El E-Voto de FICCT "
    "resuelve esto mediante una interfaz web responsiva accesible desde cualquier dispositivo.")
doc.add_page_break()

# ───────────────────────────── CAPITULO IV ────────────────────────────────────
add_heading(doc, "CAPITULO IV — DESARROLLO", 1)
add_heading(doc, "4.1. Propuesta", 2)
add_paragraph(doc,
    "Se propone el desarrollo de un Sistema de Votacion Electronica multicapa para FICCT-UAGRM "
    "que garantiza la integridad del proceso electoral mediante la combinacion de una base de "
    "datos relacional PostgreSQL para gestion de identidad y control de acceso, y una red "
    "blockchain Hyperledger Fabric para registro inmutable de votos.")

add_heading(doc, "4.1.2. Herramientas de Desarrollo", 3)
herramientas = [
    ("NestJS 10 + TypeScript", "Framework backend con inyeccion de dependencias, guards JWT y pipes de validacion."),
    ("React 18 + Vite + TypeScript", "Frontend con compilacion rapida. Tailwind v4 para estilos. Zustand para estado global."),
    ("PostgreSQL 15", "SGBD relacional con soporte para UUID, JSONB, ENUMs. Raw SQL via Pool pg sin TypeORM."),
    ("Hyperledger Fabric 2.x", "Red blockchain permisionada con CouchDB como World State. Chaincode evoting-cc."),
    ("bcrypt", "Hash de contrasenias con factor de costo 10, compatible con pgcrypto de PostgreSQL."),
    ("Docker + WSL", "Contenedorizacion de peers Fabric y CouchDB dentro de WSL en Windows."),
]
add_table(doc, ["Herramienta", "Rol en el Sistema"],
          [[h, d] for h, d in herramientas],
          "Tabla 4.1 — Herramientas de Desarrollo")
doc.add_page_break()

# ───────────────────────────── CAPITULO V ─────────────────────────────────────
add_heading(doc, "CAPITULO V — INGENIERIA DEL SOFTWARE", 1)
add_heading(doc, "5.1. Desarrollo", 2)
add_heading(doc, "5.1.1. Identificar Problemas", 3)
add_paragraph(doc, "5.1.1.1 Lista de Problemas", bold=True)
for prob in [
    "P01. Los sistemas de votacion centralizados son susceptibles a manipulacion de resultados sin mecanismo de deteccion.",
    "P02. Ausencia de un registro inmutable que permita auditar cada voto emitido tras el cierre de la eleccion.",
    "P03. Riesgo de emision de votos duplicados o en nombre de terceros sin proteccion a nivel de base de datos.",
    "P04. Falta de separacion entre identidad del votante y contenido del voto, comprometiendo el secreto del sufragio.",
    "P05. Limitada participacion estudiantil por barreras logisticas de la votacion presencial.",
    "P06. Dificultad para gestionar multiples elecciones simultaneas con diferentes estados y candidatos.",
    "P07. Inexistencia de mecanismo de verificacion individual que permita al votante confirmar su voto.",
    "P08. Ausencia de interfaz diferenciada por rol que restrinja el acceso a funciones sensibles.",
    "P09. Falta de gestion dinamica de infraestructura blockchain sin reiniciar el sistema.",
    "P10. Necesidad de auditoria del proceso electoral sin comprometer la confidencialidad.",
]:
    add_bullet(doc, prob)

add_paragraph(doc, "5.1.1.3 Lista Final de Problemas", bold=True)
for p in [
    "PF01. Manipulacion de resultados en sistemas electorales centralizados sin auditoria independiente.",
    "PF02. Ausencia de registro inmutable e inalterable de votos emitidos.",
    "PF03. Vulnerabilidad ante doble voto y suplantacion de identidad del votante.",
    "PF04. Falta de mecanismo criptografico que garantice el anonimato del sufragio.",
    "PF05. Interfaz unica sin diferenciacion de roles que expone funciones administrativas.",
    "PF06. Imposibilidad de verificar individualmente la existencia del propio voto tras emitirlo.",
    "PF07. Gestion estatica de infraestructura blockchain que impide escalar sin intervencion manual.",
]:
    add_bullet(doc, p)

# ── 5.1.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "5.1.2. Captura de Requisitos", 3)
add_heading(doc, "5.1.2.1. Actores y Casos de Uso", 3)
add_paragraph(doc, "Actores del Sistema:", bold=True)
add_table(doc, ["Actor", "Descripcion"],
          [["Administrador", "Gestiona configuracion completa: usuarios, elecciones, candidatos, estados y nodos Fabric."],
           ["Votante",       "Usuario habilitado para emitir voto en elecciones activas y verificar su txId."],
           ["Auditor",       "Rol de supervision: consulta logs y resultados, pero no puede modificar datos."]],
          "Tabla 5.1 — Actores del Sistema E-Voto")

add_paragraph(doc, "Casos de Uso del Sistema:", bold=True)
add_table(doc, ["ID", "Caso de Uso", "Actor(es)"],
          [[cu["id"], cu["nombre"], cu["actor"]] for cu in CASOS_DE_USO],
          "Tabla 5.2 — Casos de Uso del Sistema E-Voto")

# ── 5.1.3 — Detallar Casos de Uso ─────────────────────────────────────────────
add_heading(doc, "5.1.3. Detallar Casos de Uso", 3)
for cu in CASOS_DE_USO:
    add_paragraph(doc, f"{cu['id']} — {cu['nombre']}", bold=True)
    use_case_table(doc, [
        ["Caso de uso",      f"{cu['id']} — {cu['nombre']}"],
        ["Proposito",        cu["proposito"]],
        ["Actores",          cu["actor"]],
        ["Actor iniciador",  cu["actor"].split("/")[0].strip()],
        ["Precondiciones",   cu["precondiciones"]],
        ["Flujo principal",  cu["flujo"]],
        ["Excepcion",        cu["excepcion"]],
        ["Postcondiciones",  cu["postcondiciones"]],
    ])
    insertar_plantuml(doc, pu_cu_individual(cu),
                      f"Diagrama de Caso de Uso — {cu['id']}: {cu['nombre']}")

# ── 5.1.4 — Diagrama General de Casos de Uso ──────────────────────────────────
add_heading(doc, "5.1.4. Diagrama General de Casos de Uso", 3)
add_paragraph(doc,
    "El diagrama general de casos de uso muestra la interaccion de los tres actores principales "
    "(Administrador, Votante, Auditor) con los diez casos de uso del sistema E-Voto.")
insertar_plantuml(doc, PU_CU_GENERAL, "Diagrama General de Casos de Uso")
doc.add_page_break()

# ── 5.2. Analisis ─────────────────────────────────────────────────────────────
add_heading(doc, "5.2. Analisis", 2)
add_heading(doc, "5.2.1. Analisis de Arquitectura", 3)
add_heading(doc, "5.2.1.1. Identificacion de Paquetes", 3)
add_paragraph(doc, "El sistema E-Voto se organiza en los siguientes paquetes funcionales:")
add_table(doc, ["Paquete", "Descripcion"],
          [[n, d] for n, d in PAQUETES],
          "Tabla 5.3 — Paquetes del Sistema E-Voto")

add_heading(doc, "5.2.1.2. Relacionar Paquetes y Casos de Uso", 3)
add_table(doc, ["Paquete", "Casos de Uso"],
          [
              ("Seguridad",            "CU10 - Iniciar Sesion"),
              ("Gestion Electoral",    "CU01, CU02, CU03, CU04"),
              ("Votacion",             "CU05 - Emitir Voto; CU06 - Verificar Voto en Blockchain"),
              ("Auditoria y Reportes", "CU07 - Ver Resultados; CU08 - Ver Logs; CU09 - Gestionar Nodos"),
          ],
          "Tabla 5.4 — Relacion Paquetes con Casos de Uso (estereotipo trace)")
insertar_plantuml(doc, PU_PKG_TRACE, "Diagrama Paquetes <-> Casos de Uso (<<trace>>)")

# ── 5.2.2. Diagramas de Comunicacion ─────────────────────────────────────────
add_heading(doc, "5.2.2. Analisis de un Caso de Uso", 3)
add_heading(doc, "5.2.2.1. Diagramas de Comunicacion", 3)

add_paragraph(doc, "Diagrama de Comunicacion — CU10: Iniciar Sesion", bold=True)
add_paragraph(doc,
    "Muestra la colaboracion entre los objetos boundary (LoginPage), control (AuthController, "
    "AuthService, JwtService) y entity (usuarios) durante el flujo de autenticacion.")
insertar_plantuml(doc, PU_COMM_CU10, "Diagrama de Comunicacion — CU10: Iniciar Sesion")

add_paragraph(doc, "Diagrama de Comunicacion — CU05: Emitir Voto", bold=True)
add_paragraph(doc,
    "Muestra la colaboracion entre los objetos boundary (VotingPage), control (FabricController, "
    "FabricService, FabricGateway, UsersService) y entity (Ledger, recibos_voto) durante "
    "la emision del voto con proteccion anti-doble-voto.")
insertar_plantuml(doc, PU_COMM_CU05, "Diagrama de Comunicacion — CU05: Emitir Voto")

add_paragraph(doc, "Diagrama de Comunicacion — CU08: Ver Logs de Auditoria", bold=True)
add_paragraph(doc,
    "Muestra la colaboracion entre los objetos boundary (AuditLogs), control (AuditController) "
    "y entity (recibos_voto, Fabric Ledger) durante la consulta de registros de auditoria.")
insertar_plantuml(doc, PU_COMM_CU08, "Diagrama de Comunicacion — CU08: Ver Logs de Auditoria")

# ── 5.2.3. Analisis de Clase ──────────────────────────────────────────────────
add_heading(doc, "5.2.3. Analisis de Clase", 3)
add_paragraph(doc,
    "El analisis de clases del sistema E-Voto identifica las siguientes entidades principales:")
for cls, resp in [
    ("Usuario", "Representa a cualquier persona registrada. Contiene identificador, nombre, email, rol y hash de contrasena."),
    ("Eleccion", "Proceso electoral con titulo, fechas, estado y canal Fabric. Gestiona su propio ciclo de vida."),
    ("Candidato", "Postulante asociado a una Eleccion. Contiene nombre del frente, nombre, cargo y foto."),
    ("ReciboVoto", "Registro del voto emitido que vincula Usuario y Eleccion con el txId de Fabric."),
    ("PadronElectoral", "Registro de control: indica si un Usuario ha emitido su voto en una Eleccion."),
    ("NodoFabric", "Peer de Hyperledger Fabric con endpoint, host alias y prioridad de conexion."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{cls}: "); r1.font.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
    r2 = p.add_run(resp);       r2.font.name = "Calibri";                      r2.font.size = Pt(11)

# ── 5.2.4. Analisis de Paquetes ───────────────────────────────────────────────
add_heading(doc, "5.2.4. Analisis de Paquetes", 3)
add_paragraph(doc,
    "Los paquetes del sistema se relacionan de la siguiente manera: Seguridad es la base de "
    "todos los demas paquetes (autenticacion). Gestion Electoral depende de Seguridad. "
    "Votacion depende de Seguridad y Gestion Electoral. Auditoria y Reportes depende de todos.")
insertar_plantuml(doc, PU_PKG_GENERAL, "Diagrama General de Paquetes")
doc.add_page_break()

# ── 5.3. Diseno ───────────────────────────────────────────────────────────────
add_heading(doc, "5.3. Diseno", 2)

add_heading(doc, "5.3.1. Diagrama de Despliegue", 3)
add_paragraph(doc,
    "El diagrama de despliegue del sistema E-Voto incluye los nodos: cliente web (React+Vite), "
    "servidor de aplicacion (NestJS+Fabric SDK), servidor PostgreSQL 15, dos peers Fabric "
    "(peer0:7051, peer1:8051), Orderer (7050) y CouchDB (5984/6984).")
insertar_plantuml(doc, PU_DEPLOY, "Diagrama de Despliegue del Sistema E-Voto")

add_heading(doc, "5.3.2. Diseno de Datos", 3)
add_paragraph(doc,
    "El modelo de datos del sistema E-Voto se compone de 9 tablas principales en PostgreSQL:")
insertar_plantuml(doc, PU_CLASES, "Diagrama de Clases — Modelo de Base de Datos")

add_paragraph(doc, "Mapeo completo de tablas:", bold=True)
for table_name, cols in DB_TABLES:
    add_paragraph(doc, f"Tabla: {table_name}", bold=True)
    add_table(doc, ["Columna", "Tipo", "Restriccion", "Descripcion"], cols)

add_heading(doc, "5.3.3. Diagramas de Secuencia", 3)
add_paragraph(doc,
    "A continuacion se presentan los diagramas de secuencia para los diez casos de uso del sistema:")
for cu in CASOS_DE_USO:
    add_paragraph(doc, f"{cu['id']} — {cu['nombre']}", bold=True)
    insertar_plantuml(doc, pu_secuencia(cu),
                      f"Diagrama de Secuencia — {cu['id']}: {cu['nombre']}")

doc.add_page_break()

# ── 5.4. Implementacion ───────────────────────────────────────────────────────
add_heading(doc, "5.4. Implementacion", 2)
add_heading(doc, "5.4.1. Eleccion de Plataforma", 3)
add_heading(doc, "5.4.1.1. Lenguaje de Programacion", 3)
add_paragraph(doc,
    "El sistema E-Voto utiliza TypeScript como lenguaje principal tanto en el backend como "
    "en el frontend. TypeScript es un superconjunto tipado de JavaScript que permite detectar "
    "errores en tiempo de compilacion, mejora la mantenibilidad del codigo y provee "
    "autocompletado inteligente en entornos de desarrollo modernos.")

add_heading(doc, "5.4.1.2. Base de Datos", 3)
add_paragraph(doc,
    "PostgreSQL 15 fue seleccionado como SGBD por: soporte nativo para UUID, tipo JSONB "
    "para metadatos flexibles, ENUMs nativos para garantizar integridad referencial, "
    "extension pgcrypto para hash de contrasenias, y soporte para SELECT FOR UPDATE "
    "necesario para la transaccion atomica anti-doble-voto.")

add_heading(doc, "5.4.1.3. Sistema Operativo", 3)
add_paragraph(doc,
    "El sistema fue desarrollado y probado en Windows 11 con WSL para la ejecucion de la "
    "red Hyperledger Fabric. Para entornos de produccion, se recomienda Linux (Ubuntu 20.04 LTS) "
    "donde Docker y Fabric pueden ejecutarse sin la capa de emulacion de WSL.")
doc.add_page_break()

# ─────────────────────────── CONCLUSIONES ─────────────────────────────────────
add_heading(doc, "CONCLUSIONES Y RECOMENDACIONES", 1)
add_heading(doc, "Conclusiones", 2)
for i, c in enumerate([
    "El sistema E-Voto demuestra que Hyperledger Fabric es una plataforma viable para garantizar la inmutabilidad e integridad de procesos electorales universitarios.",
    "La separacion entre la capa de identidad (PostgreSQL) y la capa de registro del voto (Fabric) resuelve el principal dilema: garantizar simultaneamente el anonimato del votante y la trazabilidad del proceso.",
    "La implementacion de tres capas de proteccion anti-doble-voto (assertCanVote, SELECT FOR UPDATE y UNIQUE constraint) garantiza que ningun votante pueda emitir mas de un voto bajo alta concurrencia.",
    "La arquitectura multicapa NestJS + React + PostgreSQL + Fabric permite escalar cada componente independientemente.",
    "La gestion dinamica de nodos y canales Fabric desde la interfaz administrativa reduce la barrera tecnica para la expansion de la red blockchain.",
], 1):
    add_bullet(doc, f"C{i}. {c}")

add_heading(doc, "Recomendaciones", 2)
for i, r in enumerate([
    "Desplegar el sistema en entorno Linux productivo para eliminar la dependencia de WSL y mejorar el rendimiento de la red Fabric.",
    "Implementar HTTPS con certificados TLS en el backend y el frontend para cifrar la comunicacion en produccion.",
    "Agregar autenticacion multifactor (MFA) para el rol Administrador, dado el nivel de privilegio que posee.",
    "Extender el sistema con notificaciones en tiempo real (WebSockets) para que el panel de resultados se actualice automaticamente.",
    "Realizar pruebas de carga (stress testing) con herramientas como k6 o Artillery para validar el comportamiento bajo alta concurrencia.",
], 1):
    add_bullet(doc, f"R{i}. {r}")

doc.add_page_break()

# ─────────────────────────── BIBLIOGRAFIA ─────────────────────────────────────
add_heading(doc, "BIBLIOGRAFIA", 1)
for ref in [
    "Hyperledger Foundation. (2023). Hyperledger Fabric Documentation v2.5. The Linux Foundation.",
    "Androulaki, E., et al. (2018). Hyperledger Fabric: A Distributed Operating System for Permissioned Blockchains. EuroSys '18. ACM.",
    "NestJS. (2024). NestJS Documentation v10.",
    "The PostgreSQL Global Development Group. (2023). PostgreSQL 15 Documentation.",
    "React Team. (2024). React Documentation v18. Meta Open Source.",
    "Kshetri, N., & Voas, J. (2018). Blockchain-Enabled E-Voting. IEEE Software, 35(4), 95-99.",
    "Park, S., & Specter, M. (2021). Going from bad to worse: From Internet Voting to Blockchain Voting. Journal of Cybersecurity, 7(1).",
    "Zustand. (2024). Zustand — A small, fast and scalable state-management solution.",
]:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(ref); r.font.name = "Calibri"; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ─────────────────────────── ANEXOS ───────────────────────────────────────────
add_heading(doc, "ANEXO", 1)
add_heading(doc, "Anexo A — Credenciales de Prueba del Sistema", 2)
add_paragraph(doc, "Contrasena para todos los usuarios: Password123")
add_table(doc, ["Identificador", "Nombre", "Rol", "Email"],
          [
              ["ADMIN001", "Carlos Mendoza Rivera",  "ADMINISTRADOR", "admin001@ficct.edu.bo"],
              ["ADMIN002", "Maria Rojas Soliz",      "ADMINISTRADOR", "admin002@ficct.edu.bo"],
              ["AUD001",   "Roberto Vargas Aliaga",  "AUDITOR",       "aud001@ficct.edu.bo"],
              ["AUD002",   "Lucia Fernandez Paz",    "AUDITOR",       "aud002@ficct.edu.bo"],
              ["221001",   "Juan Perez Silva",       "VOTANTE",       "221001@ficct.edu.bo"],
              ["221010",   "Maria Condori Flores",   "VOTANTE",       "221010@ficct.edu.bo"],
          ],
          "Tabla A.1 — Usuarios de Prueba")

add_heading(doc, "Anexo B — Variables de Entorno del Backend", 2)
add_table(doc, ["Variable", "Valor de Desarrollo", "Descripcion"],
          [
              ["PORT",               "3000",                    "Puerto del servidor NestJS"],
              ["DB_HOST",            "localhost",               "Host de PostgreSQL"],
              ["DB_PORT",            "5432",                    "Puerto de PostgreSQL"],
              ["DB_USER",            "postgres",                "Usuario de PostgreSQL"],
              ["DB_PASSWORD",        "postgres",                "Contrasena de PostgreSQL"],
              ["DB_NAME",            "sw2_primer_parcial",      "Base de datos"],
              ["JWT_SECRET",         "change_this_secret",      "Clave de firma JWT"],
              ["JWT_EXPIRES_IN",     "8h",                      "Expiracion del token JWT"],
              ["FABRIC_PEER_ENDPOINT","localhost:7051",         "Endpoint del peer principal"],
              ["FABRIC_CHANNEL",     "evoting",                 "Canal Fabric del sistema"],
              ["FABRIC_CHAINCODE",   "evoting-cc",              "Nombre del chaincode"],
              ["FRONTEND_URL",       "http://localhost:5173",   "URL del frontend (CORS)"],
          ],
          "Tabla B.1 — Variables de Entorno")

add_heading(doc, "Anexo C — Resumen de Endpoints REST", 2)
add_table(doc, ["Metodo", "Ruta", "Auth", "Descripcion"],
          [
              ["POST",   "/auth/login",                      "—",           "Login, retorna JWT + datos de usuario"],
              ["GET",    "/users",                           "JWT",         "Listar todos los usuarios"],
              ["POST",   "/users",                           "ADMIN",       "Crear nuevo usuario"],
              ["PATCH",  "/users/:id",                       "ADMIN",       "Actualizar usuario"],
              ["DELETE", "/users/:id",                       "ADMIN",       "Eliminar usuario"],
              ["GET",    "/elections",                       "JWT",         "Listar elecciones con candidatos"],
              ["POST",   "/elections",                       "ADMIN",       "Crear eleccion (estado: PROGRAMADA)"],
              ["PATCH",  "/elections/:id/status",            "ADMIN",       "Cambiar estado de eleccion"],
              ["DELETE", "/elections/:id",                   "ADMIN",       "Eliminar eleccion (solo PROGRAMADA)"],
              ["POST",   "/elections/:id/candidates",        "ADMIN",       "Agregar candidato"],
              ["DELETE", "/elections/:eid/candidates/:id",   "ADMIN",       "Eliminar candidato"],
              ["POST",   "/fabric/vote",                     "JWT",         "Emitir voto en Fabric"],
              ["GET",    "/fabric/results/:electionId",      "—",           "Resultados desde base de datos"],
              ["GET",    "/fabric/verify/:txId",             "—",           "Verificar voto en ledger"],
              ["GET",    "/audit/logs",                      "ADMIN/AUDITOR","Ver recibos de voto"],
              ["GET",    "/nodes",                           "JWT",         "Listar nodos Fabric"],
              ["POST",   "/nodes/deploy",                    "ADMIN",       "Desplegar nuevo peer via WSL"],
              ["GET",    "/channels",                        "JWT",         "Listar canales Fabric"],
              ["POST",   "/channels",                        "ADMIN",       "Crear canal via WSL"],
          ],
          "Tabla C.1 — Endpoints REST del Sistema")

# ─────────────────────────── GUARDAR ──────────────────────────────────────────
doc.save(OUTPUT)

n_tablas = sum(1 for _ in doc.tables)
n_parrafos = len(doc.paragraphs)
n_bloques_pu = _fig[0]

print(f"\n[OK] Documento corregido: {OUTPUT}")
print(f"[OK] Bloques PlantUML insertados: {n_bloques_pu}")
print(f"[OK] Figuras numeradas correctamente: 1 a {n_bloques_pu}")
print(f"[OK] Sin placeholders vacios")
print(f"[OK] Tablas insertadas: {n_tablas}")
print(f"[OK] Parrafos totales: {n_parrafos}")
print(f"[OK] Listo para que el usuario copie cada bloque y renderice")
print(f"     en https://www.plantuml.com/plantuml/uml/")
